import streamlit as st
import requests
from typing import List
import time
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import io

# Page configuration
st.set_page_config(
    page_title="Resume Comparison Tool",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS with dark theme support
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(90deg, #1f77b4 0%, #2ca02c 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 20px 0;
    }
    
    /* Result boxes with dark theme support */
    .duplicate-box {
        background-color: rgba(255, 68, 68, 0.15);
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #ff4444;
        margin: 15px 0;
        color: inherit;
    }
    
    .similar-box {
        background-color: rgba(255, 193, 7, 0.15);
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #ffc107;
        margin: 15px 0;
        color: inherit;
    }
    
    .success-box {
        background-color: rgba(40, 167, 69, 0.15);
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #28a745;
        margin: 15px 0;
        color: inherit;
    }
    
    .info-box {
        background-color: rgba(31, 119, 180, 0.15);
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #1f77b4;
        margin: 15px 0;
        color: inherit;
    }
    
    /* File card styling */
    .file-card {
        background-color: rgba(128, 128, 128, 0.1);
        padding: 12px;
        border-radius: 8px;
        margin: 8px 0;
        border: 1px solid rgba(128, 128, 128, 0.2);
    }
    
    /* Metric cards */
    .metric-card {
        background: linear-gradient(135deg, rgba(31, 119, 180, 0.2) 0%, rgba(44, 160, 44, 0.2) 100%);
        padding: 15px;
        border-radius: 10px;
        text-align: center;
        border: 1px solid rgba(128, 128, 128, 0.2);
    }
    
    /* Strong text visibility */
    .duplicate-box strong, .similar-box strong, .success-box strong, .info-box strong {
        font-weight: 700;
        font-size: 1.1em;
    }
    
    /* Button hover effects */
    .stButton>button {
        transition: all 0.3s ease;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15);
    }
    </style>
""", unsafe_allow_html=True)

# API Configuration
API_URL = "http://localhost:8000"

def check_api_health():
    """Check if the API is running"""
    try:
        response = requests.get(f"{API_URL}/health", timeout=2)
        return response.status_code == 200
    except:
        return False

def compare_resumes_api(files: List, threshold: float):
    """Call the FastAPI backend to compare resumes"""
    try:
        files_data = []
        for uploaded_file in files:
            files_data.append(
                ('files', (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type))
            )
        
        response = requests.post(
            f"{API_URL}/compare-resumes",
            files=files_data,
            params={"similarity_threshold": threshold}
        )
        
        if response.status_code == 200:
            return response.json(), None
        else:
            return None, f"Error: {response.json().get('detail', 'Unknown error')}"
    except requests.exceptions.ConnectionError:
        return None, "Cannot connect to API. Please ensure the FastAPI backend is running."
    except Exception as e:
        return None, f"Error: {str(e)}"

def generate_txt_report(result):
    """Generate TXT report"""
    report = []
    report.append("=" * 80)
    report.append("RESUME COMPARISON REPORT")
    report.append("=" * 80)
    report.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append("")
    
    # Summary
    summary = result['summary']
    report.append("SUMMARY")
    report.append("-" * 80)
    report.append(f"Total Files Analyzed: {summary['total_files']}")
    report.append(f"Unique Files: {summary['unique_files']}")
    report.append(f"Exact Duplicate Groups: {summary['exact_duplicate_groups']}")
    report.append(f"Similar Pairs Found: {summary['similar_pairs_found']}")
    report.append(f"Similarity Threshold: {summary['similarity_threshold']}%")
    report.append("")
    
    # Exact Duplicates
    report.append("EXACT DUPLICATES")
    report.append("-" * 80)
    if result['exact_duplicates']:
        for i, group in enumerate(result['exact_duplicates'], 1):
            report.append(f"\nDuplicate Group {i}:")
            for file in group:
                report.append(f"  - {file}")
    else:
        report.append("No exact duplicates found.")
    report.append("")
    
    # Similar Pairs
    report.append("SIMILAR RESUMES")
    report.append("-" * 80)
    if result['similar_pairs']:
        for pair in result['similar_pairs']:
            report.append(f"\nSimilarity: {pair['similarity']}%")
            report.append(f"  File 1: {pair['file1']}")
            report.append(f"  File 2: {pair['file2']}")
    else:
        report.append(f"No similar resumes found above {summary['similarity_threshold']}% threshold.")
    
    report.append("")
    report.append("=" * 80)
    report.append("END OF REPORT")
    report.append("=" * 80)
    
    return "\n".join(report)

def generate_docx_report(result):
    """Generate DOCX report"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Resume Comparison Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Timestamp
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Summary Section
    doc.add_heading('Summary', level=1)
    summary = result['summary']
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Light Grid Accent 1'
    
    summary_data = [
        ('Total Files Analyzed', str(summary['total_files'])),
        ('Unique Files', str(summary['unique_files'])),
        ('Exact Duplicate Groups', str(summary['exact_duplicate_groups'])),
        ('Similar Pairs Found', str(summary['similar_pairs_found'])),
        ('Similarity Threshold', f"{summary['similarity_threshold']}%")
    ]
    
    for i, (key, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = key
        summary_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Exact Duplicates
    doc.add_heading('Exact Duplicates', level=1)
    if result['exact_duplicates']:
        for i, group in enumerate(result['exact_duplicates'], 1):
            doc.add_heading(f'Duplicate Group {i}', level=2)
            for file in group:
                doc.add_paragraph(file, style='List Bullet')
    else:
        doc.add_paragraph('✓ No exact duplicates found.')
    
    doc.add_paragraph()
    
    # Similar Resumes
    doc.add_heading('Similar Resumes', level=1)
    if result['similar_pairs']:
        for pair in result['similar_pairs']:
            p = doc.add_paragraph()
            p.add_run(f"Similarity: {pair['similarity']}%\n").bold = True
            p.add_run(f"File 1: {pair['file1']}\n")
            p.add_run(f"File 2: {pair['file2']}")
            doc.add_paragraph()
    else:
        doc.add_paragraph(f"✓ No similar resumes found above {summary['similarity_threshold']}% threshold.")
    
    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

def generate_pdf_report(result):
    """Generate PDF report"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1f77b4'),
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#2ca02c'),
        spaceAfter=12
    )
    
    # Title
    story.append(Paragraph("Resume Comparison Report", title_style))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    # Summary
    story.append(Paragraph("Summary", heading_style))
    summary = result['summary']
    summary_data = [
        ['Metric', 'Value'],
        ['Total Files Analyzed', str(summary['total_files'])],
        ['Unique Files', str(summary['unique_files'])],
        ['Exact Duplicate Groups', str(summary['exact_duplicate_groups'])],
        ['Similar Pairs Found', str(summary['similar_pairs_found'])],
        ['Similarity Threshold', f"{summary['similarity_threshold']}%"]
    ]
    
    summary_table = Table(summary_data, colWidths=[3.5*inch, 2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f77b4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3*inch))
    
    # Exact Duplicates
    story.append(Paragraph("Exact Duplicates", heading_style))
    if result['exact_duplicates']:
        for i, group in enumerate(result['exact_duplicates'], 1):
            story.append(Paragraph(f"<b>Duplicate Group {i}:</b>", styles['Normal']))
            for file in group:
                story.append(Paragraph(f"• {file}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    else:
        story.append(Paragraph("✓ No exact duplicates found.", styles['Normal']))
    story.append(Spacer(1, 0.2*inch))
    
    # Similar Resumes
    story.append(Paragraph("Similar Resumes", heading_style))
    if result['similar_pairs']:
        for pair in result['similar_pairs']:
            story.append(Paragraph(f"<b>Similarity: {pair['similarity']}%</b>", styles['Normal']))
            story.append(Paragraph(f"File 1: {pair['file1']}", styles['Normal']))
            story.append(Paragraph(f"File 2: {pair['file2']}", styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
    else:
        story.append(Paragraph(f"✓ No similar resumes found above {summary['similarity_threshold']}% threshold.", styles['Normal']))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def main():
    # Header
    st.markdown('<p class="main-header">📄 Resume Comparison Tool</p>', unsafe_allow_html=True)
    
    # Add tabs for better organization
    tab1, tab2, tab3 = st.tabs(["🔍 Compare", "📊 Results", "ℹ️ About"])
    
    with tab1:
        st.markdown("---")
        
        # Sidebar
        with st.sidebar:
            st.header("⚙️ Settings")
            
            # API Status
            api_status = check_api_health()
            if api_status:
                st.success("✅ API Connected")
            else:
                st.error("❌ API Disconnected")
                st.warning("Please start the FastAPI backend:\n```bash\npython app.py```")
            
            st.markdown("---")
            
            # Similarity threshold
            similarity_threshold = st.slider(
                "Similarity Threshold (%)",
                min_value=50,
                max_value=100,
                value=95,
                step=5,
                help="Set the minimum similarity percentage to flag resumes as similar"
            )
            
            st.markdown("---")
            
            # Statistics
            if 'result' in st.session_state and st.session_state.result:
                st.subheader("📈 Quick Stats")
                result = st.session_state.result
                summary = result['summary']
                
                st.metric("Total Files", summary['total_files'], delta=None)
                st.metric("Unique Files", summary['unique_files'], delta=None)
                st.metric("Duplicates", summary['exact_duplicate_groups'], delta=None, delta_color="inverse")
                st.metric("Similar Pairs", summary['similar_pairs_found'], delta=None)
            
            st.markdown("---")
            
            # Instructions
            st.subheader("📖 How to Use")
            st.markdown("""
            1. Upload 2+ resumes
            2. Adjust threshold
            3. Click 'Compare'
            4. View & download results
            
            **Formats:**
            - PDF (.pdf)
            - Word (.docx)
            - Text (.txt)
            """)
        
        # Main content
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.subheader("📤 Upload Resumes")
            uploaded_files = st.file_uploader(
                "Choose resume files",
                type=['pdf', 'docx', 'txt'],
                accept_multiple_files=True,
                help="Upload multiple resumes to compare"
            )
            
            if uploaded_files:
                st.markdown(f"""
                <div class="info-box">
                    <strong>📁 {len(uploaded_files)} file(s) uploaded</strong>
                </div>
                """, unsafe_allow_html=True)
                
                with st.expander("📋 View uploaded files", expanded=False):
                    for i, file in enumerate(uploaded_files, 1):
                        file_size_kb = file.size / 1024
                        file_type = file.type.split('/')[-1].upper()
                        st.markdown(f"""
                        <div class="file-card">
                            <strong>{i}.</strong> {file.name}<br>
                            <small>Size: {file_size_kb:.2f} KB | Type: {file_type}</small>
                        </div>
                        """, unsafe_allow_html=True)
        
        with col2:
            st.subheader("🔍 Analysis")
            st.markdown(f"""
            <div class="metric-card">
                <h2>{len(uploaded_files) if uploaded_files else 0}</h2>
                <p>Files Uploaded</p>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown(f"""
            <div class="metric-card">
                <h2>{similarity_threshold}%</h2>
                <p>Threshold</p>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Compare button
        if uploaded_files and len(uploaded_files) >= 2:
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                if st.button("🚀 Compare Resumes", type="primary", use_container_width=True):
                    with st.spinner("🔄 Analyzing resumes..."):
                        # Progress bar
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        for i in range(100):
                            time.sleep(0.01)
                            progress_bar.progress(i + 1)
                            if i < 30:
                                status_text.text("📄 Reading files...")
                            elif i < 60:
                                status_text.text("🔍 Detecting duplicates...")
                            elif i < 90:
                                status_text.text("📊 Calculating similarities...")
                            else:
                                status_text.text("✅ Finalizing results...")
                        
                        # Call API
                        result, error = compare_resumes_api(uploaded_files, similarity_threshold / 100)
                        
                        progress_bar.empty()
                        status_text.empty()
                        
                        if error:
                            st.error(error)
                        else:
                            st.session_state.result = result
                            st.success("✅ Analysis complete! Check the Results tab.")
                            st.balloons()
        else:
            st.info("⚠️ Please upload at least 2 resume files to compare")
    
    with tab2:
        if 'result' in st.session_state and st.session_state.result:
            display_results(st.session_state.result)
        else:
            st.info("👈 Upload files and run comparison to see results here")
    
    with tab3:
        display_about()

def display_results(result):
    """Display comparison results"""
    st.header("📊 Comparison Results")
    
    # Summary metrics
    summary = result['summary']
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("📁 Total Files", summary['total_files'])
    with col2:
        st.metric("✅ Unique Files", summary['unique_files'])
    with col3:
        st.metric("🔴 Duplicate Groups", summary['exact_duplicate_groups'])
    with col4:
        st.metric("🟡 Similar Pairs", summary['similar_pairs_found'])
    
    st.markdown("---")
    
    # Download buttons
    st.subheader("📥 Download Report")
    col1, col2, col3, col4 = st.columns(4)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    with col1:
        json_data = json.dumps(result, indent=2)
        st.download_button(
            label="📄 JSON",
            data=json_data,
            file_name=f"resume_report_{timestamp}.json",
            mime="application/json",
            use_container_width=True
        )
    
    with col2:
        txt_data = generate_txt_report(result)
        st.download_button(
            label="📝 TXT",
            data=txt_data,
            file_name=f"resume_report_{timestamp}.txt",
            mime="text/plain",
            use_container_width=True
        )
    
    with col3:
        docx_data = generate_docx_report(result)
        st.download_button(
            label="📘 DOCX",
            data=docx_data,
            file_name=f"resume_report_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col4:
        pdf_data = generate_pdf_report(result)
        st.download_button(
            label="📕 PDF",
            data=pdf_data,
            file_name=f"resume_report_{timestamp}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
    
    st.markdown("---")
    
    # Exact Duplicates
    st.subheader("🔴 Exact Duplicates")
    exact_duplicates = result['exact_duplicates']
    
    if exact_duplicates:
        for i, group in enumerate(exact_duplicates, 1):
            files_list = '<br>'.join([f'<strong>•</strong> {file}' for file in group])
            st.markdown(f"""
            <div class="duplicate-box">
                <strong>🔴 Duplicate Group {i}</strong><br><br>
                <em>The following files are IDENTICAL (byte-for-byte):</em><br><br>
                {files_list}
            </div>
            """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div class="success-box">
            <strong>✅ No exact duplicates found!</strong><br>
            All files are unique.
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Similar Pairs
    st.subheader("🟡 Similar Resumes")
    similar_pairs = result['similar_pairs']
    
    if similar_pairs:
        for pair in similar_pairs:
            similarity_color = "🔴" if pair['similarity'] >= 98 else "🟡" if pair['similarity'] >= 90 else "🟢"
            st.markdown(f"""
            <div class="similar-box">
                <strong>{similarity_color} Similarity: {pair['similarity']}%</strong><br><br>
                📄 <strong>File 1:</strong> {pair['file1']}<br>
                📄 <strong>File 2:</strong> {pair['file2']}
            </div>
            """, unsafe_allow_html=True)
    else:
        st.markdown(f"""
        <div class="success-box">
            <strong>✅ No similar resumes found!</strong><br>
            No pairs exceeded the {summary['similarity_threshold']}% similarity threshold.
        </div>
        """, unsafe_allow_html=True)

def display_about():
    """Display about information"""
    st.header("ℹ️ About Resume Comparison Tool")
    
    st.markdown("""
    <div class="info-box">
        <h3>🎯 Purpose</h3>
        <p>This tool helps you identify duplicate and similar resumes efficiently using advanced text analysis algorithms.</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-box">
        <h3>✨ Features</h3>
        <ul>
            <li><strong>Exact Duplicate Detection:</strong> Identifies identical files using MD5 hash comparison</li>
            <li><strong>Similarity Analysis:</strong> Detects content similarities using SequenceMatcher algorithm</li>
            <li><strong>Multi-Format Support:</strong> Works with PDF, DOCX, and TXT files</li>
            <li><strong>Adjustable Threshold:</strong> Customize sensitivity for similarity detection</li>
            <li><strong>Multiple Export Formats:</strong> Download reports in JSON, TXT, DOCX, or PDF</li>
            <li><strong>Dark Mode Compatible:</strong> Beautiful UI in both light and dark themes</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-box">
        <h3>📊 How Similarity Works</h3>
        <p>The tool uses the <strong>SequenceMatcher</strong> algorithm which:</p>
        <ul>
            <li>Compares text content character by character</li>
            <li>Calculates a ratio between 0-100%</li>
            <li>100% = Identical content</li>
            <li>95-99% = Very similar (minor differences)</li>
            <li>80-94% = Moderately similar</li>
            <li>Below 80% = Different resumes</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-box">
        <h3>💡 Tips for Best Results</h3>
        <ul>
            <li>Use consistent file formats for better comparison</li>
            <li>Adjust the threshold based on your needs (higher = stricter)</li>
            <li>For exact duplicates, the threshold doesn't matter</li>
            <li>Large PDFs may take longer to process</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-box">
        <h3>🔧 Technology Stack</h3>
        <ul>
            <li><strong>Frontend:</strong> Streamlit</li>
            <li><strong>Backend:</strong> FastAPI</li>
            <li><strong>PDF Processing:</strong> PyPDF2</li>
            <li><strong>DOCX Processing:</strong> python-docx</li>
            <li><strong>Report Generation:</strong> ReportLab, python-docx</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    # Initialize session state
    if 'result' not in st.session_state:
        st.session_state.result = None
    
    main()